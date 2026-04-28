import sys
from docx import Document
from docx.shared import Pt
import re

def convert_md_to_docx(md_path, docx_path):
    document = Document()
    
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except Exception as e:
        print(f"Error reading {md_path}: {e}")
        sys.exit(1)

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('### '):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith('---'):
            pass  # omit horizontal rules or maybe add page breaks if you like
        elif line.startswith('* '):
            document.add_paragraph(line[2:], style='List Bullet')
        else:
            # We can parse basic bold texting using double asterisk
            # However python-docx adds markdown quite manually. At least paragraph:
            p = document.add_paragraph()
            
            # Simple markdown bold parser **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    try:
        document.save(docx_path)
        print(f"Successfully saved {docx_path}")
    except Exception as e:
        print(f"Error saving {docx_path}: {e}")
        sys.exit(1)

if __name__ == '__main__':
    in_file = r'C:\Users\venky\.gemini\antigravity\brain\62a1f058-a3e2-4438-9d3c-4be2a71302a2\artifacts\Creator_Link_Project_Report.md'
    out_file = r'c:\Users\venky\OneDrive\Desktop\Creator Link\Creator_Link_Project_Report.docx'
    convert_md_to_docx(in_file, out_file)
